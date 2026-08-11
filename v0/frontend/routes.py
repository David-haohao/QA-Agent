# ============================================================
# 前端路由 — FastAPI路由定义
# 页面渲染、SSE流式问答、反馈提交、会话管理、文档预览
# ============================================================

import os
import uuid
import json
from urllib.parse import quote

from fastapi import APIRouter, Request, Form, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse, FileResponse
from fastapi.templating import Jinja2Templates

# 创建模板引擎（使用绝对路径避免目录切换时找不到模板）
_TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
templates = Jinja2Templates(directory=_TEMPLATES_DIR)

# 创建路由
router = APIRouter()


def create_frontend_routes(qa_service):
    """
    创建前端的FastAPI路由
    qa_service: 问答服务实例(封装了预处理+检索+生成+后处理的完整链路)
    """

    @router.get("/", response_class=HTMLResponse)
    async def index(request: Request):
        """主聊天页面"""
        session_id = request.cookies.get("session_id", f"sess_{uuid.uuid4().hex[:16]}")
        response = templates.TemplateResponse(
            request=request,
            name="index.html",
            context={"request": request, "session_id": session_id},
        )
        if not request.cookies.get("session_id"):
            response.set_cookie(key="session_id", value=session_id)
        return response

    @router.post("/api/chat")
    async def chat(request: Request):
        """
        问答接口 — 接收用户问题，返回SSE流式响应
        """
        form_data = await request.form()
        query = form_data.get("query", "").strip()
        session_id = form_data.get("session_id", "")
        if not session_id:
            session_id = request.cookies.get("session_id", f"sess_{uuid.uuid4().hex[:16]}")

        if not query:
            return JSONResponse({"error": "查询不能为空"}, status_code=400)

        # 返回SSE流
        async def event_stream():
            async for event, data in qa_service.process_query_stream(
                query=query,
                session_id=session_id,
            ):
                yield f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"

        response = StreamingResponse(
            event_stream(),
            media_type="text/event-stream",
        )
        response.set_cookie(key="session_id", value=session_id)
        return response

    @router.post("/api/feedback")
    async def feedback(request: Request):
        """提交反馈（点赞/点踩/纠错）"""
        form_data = await request.form()
        query_id = form_data.get("query_id", "")
        rating = int(form_data.get("rating", 0))
        comment = form_data.get("comment", "")
        session_id = request.cookies.get("session_id", "")
        correction_type = form_data.get("correction_type", "")
        correction_text = form_data.get("correction_text", "")
        query = form_data.get("query", "")
        answer = form_data.get("answer", "")
        source_docs_json = form_data.get("source_docs", "[]")

        try:
            source_docs = json.loads(source_docs_json)
        except (json.JSONDecodeError, TypeError):
            source_docs = []

        qa_service.record_feedback(
            session_id=session_id,
            query_id=query_id,
            rating=rating,
            comment=comment,
            correction_type=correction_type,
            correction_text=correction_text,
            query=query,
            answer=answer,
            source_docs=source_docs,
        )
        return JSONResponse({"status": "ok", "feedback_id": f"fb_{uuid.uuid4().hex[:8]}"})

    @router.post("/api/analytics/event")
    async def analytics_event(request: Request):
        """前端埋点事件"""
        form_data = await request.form()
        event_type = form_data.get("event_type", "")
        session_id = request.cookies.get("session_id", "")
        source_query_id = form_data.get("source_query_id", "")
        metadata = form_data.get("metadata", "{}")

        qa_service.record_analytics_event(
            session_id=session_id,
            event_type=event_type,
            source_query_id=source_query_id,
            metadata=json.loads(metadata) if metadata else {},
        )
        return JSONResponse({"status": "ok"})

    @router.post("/api/session")
    async def create_session(request: Request):
        """创建新会话"""
        form_data = await request.form()
        user_name = form_data.get("user_name", "")
        session_id = f"sess_{uuid.uuid4().hex[:16]}"

        qa_service.create_session(session_id, user_name)
        response = JSONResponse({
            "session_id": session_id,
            "created_at": __import__("datetime").datetime.now().isoformat(),
        })
        response.set_cookie(key="session_id", value=session_id)
        return response

    @router.get("/api/kb/overview")
    async def kb_overview():
        """知识库概览"""
        info = qa_service.get_kb_overview()
        return JSONResponse(info)

    @router.get("/api/suggestions")
    async def suggestions():
        """获取建议问题列表（高分问答对 + 热门问题）"""
        questions = qa_service.get_suggestion_questions()
        return JSONResponse({"questions": questions})

    @router.get("/api/source-detail")
    async def source_detail(filename: str = ""):
        """根据文档名获取文档详情（检查是否可查看全文）"""
        content, matched = qa_service.get_document_full_text(filename)
        return JSONResponse({
            "name": filename,
            "matched_name": matched,
            "found": content is not None,
            "view_url": f"/kb/view/{quote(filename, safe='')}" if content else None,
        })

    @router.get("/kb/view/{filename:path}", response_class=HTMLResponse)
    async def view_document(request: Request, filename: str):
        """查看文档预览——点击来源超链接后立即展示文档内容（支持模糊匹配）
        - PDF: 直接提供PDF文件（浏览器内嵌预览）
        - Word/Excel: 提供预转换的HTML预览页面
        所有文档在 build-kb 时已预转换为可预览格式，确保点击即看"""
        import re

        root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        kb_data_dir = qa_service.config.kb_config.get("kb_data_dir", "./kb_data")
        if not os.path.isabs(kb_data_dir):
            kb_data_dir = os.path.join(root_dir, kb_data_dir)
        docs_dir = qa_service.config.kb_config.get("documents_dir", "./documents")
        if not os.path.isabs(docs_dir):
            docs_dir = os.path.join(root_dir, docs_dir)

        # 解析真实文件名
        resolved_name = filename

        def try_resolve(name):
            """尝试解析文件名"""
            # 精确匹配
            doc_path = os.path.join(docs_dir, name)
            if os.path.exists(doc_path):
                return name

            # 模糊匹配
            matched = qa_service._match_doc_in_kb(name)
            if matched:
                m = matched[0]
                if os.path.exists(os.path.join(docs_dir, m)):
                    return m

            # 扫描匹配
            if os.path.isdir(docs_dir):
                clean = os.path.splitext(name)[0]
                clean = re.sub(r'^\d+[_\-\s]*', '', clean).strip()
                for fname in sorted(os.listdir(docs_dir)):
                    fbase = os.path.splitext(fname)[0]
                    if clean in fbase or fbase in clean or name in fname:
                        if os.path.exists(os.path.join(docs_dir, fname)):
                            return fname
            return None

        resolved_name = try_resolve(filename)

        if not resolved_name:
            return HTMLResponse(
                f"<h2>文档未找到</h2><p>查询名称: {escape_html(filename)}</p>"
                f"<p>该名称与知识库中任何文档都不匹配，请确认知识库已构建（python run.py build-kb）</p>"
                f"<p><a href='/'>← 返回对话</a></p>",
                status_code=404)

        doc_path = os.path.join(docs_dir, resolved_name)
        fname_lower = resolved_name.lower()

        # PDF: 直接返回PDF文件（浏览器内嵌预览，不含中文Content-Disposition）
        if fname_lower.endswith(".pdf"):
            return FileResponse(
                path=doc_path,
                media_type="application/pdf",
            )

        # Word / Excel: 返回预转换的HTML预览
        # 预转换的HTML文件保存在 kb_data/docs_html/ 目录
        docs_html_dir = os.path.join(kb_data_dir, "docs_html")
        html_filename = resolved_name + ".html"
        html_path = os.path.join(docs_html_dir, html_filename)

        if os.path.exists(html_path):
            with open(html_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            return HTMLResponse(html_content)

        # 如果预转换HTML不存在，现场生成
        html_content = _generate_doc_html_preview(doc_path, resolved_name)
        if html_content:
            # 保存以备下次使用
            os.makedirs(docs_html_dir, exist_ok=True)
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            return HTMLResponse(html_content)

        # 完全无法预览时，触发下载
        return FileResponse(
            path=doc_path,
            media_type="application/octet-stream",
            filename=resolved_name,
        )

    return router


def escape_html(text: str) -> str:
    """转义HTML特殊字符"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _generate_doc_html_preview(doc_path: str, filename: str) -> str:
    """现场生成文档的HTML预览页面
    - Word (.docx/.doc): 提取段落文本渲染为HTML
    - Excel (.xlsx/.xls): 提取表格数据渲染为HTML表格
    返回HTML字符串，失败返回None"""
    fname_lower = filename.lower()

    try:
        if fname_lower.endswith(".docx"):
            return _docx_to_html(doc_path, filename)
        elif fname_lower.endswith(".doc"):
            return _doc_to_html(doc_path, filename)
        elif fname_lower.endswith((".xlsx", ".xls")):
            return _excel_to_html(doc_path, filename)
    except Exception as e:
        print(f"生成HTML预览失败: {filename}, 错误: {e}")
    return None


def _docx_to_html(doc_path: str, filename: str) -> str:
    """将.docx文档转换为HTML预览"""
    from docx import Document

    doc = Document(doc_path)
    safe_name = escape_html(filename)

    parts = [_doc_html_header(safe_name)]

    for para in doc.paragraphs:
        text = escape_html(para.text.strip())
        if not text:
            parts.append("<br>")
            continue

        # 根据样式判断标题级别
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading 1") or style_name == "Title":
            parts.append(f'<h2 class="doc-h2">{text}</h2>')
        elif style_name.startswith("Heading 2"):
            parts.append(f'<h3 class="doc-h3">{text}</h3>')
        elif style_name.startswith("Heading"):
            parts.append(f'<h4 class="doc-h4">{text}</h4>')
        else:
            # 检测加粗
            is_bold = any(run.bold for run in para.runs if run.bold)
            if is_bold and len(text) < 80:
                parts.append(f'<p class="doc-p"><strong>{text}</strong></p>')
            else:
                parts.append(f'<p class="doc-p">{text}</p>')

    # 表格
    for table in doc.tables:
        parts.append('<table class="doc-table">')
        for row in table.rows:
            parts.append('<tr>')
            for cell in row.cells:
                cell_text = escape_html(cell.text.strip())
                parts.append(f'<td>{cell_text}</td>')
            parts.append('</tr>')
        parts.append('</table>')

    parts.append(_doc_html_footer())
    return "\n".join(parts)


def _doc_to_html(doc_path: str, filename: str) -> str:
    """将旧版.doc文档转换为HTML预览（通过textutil提取文本）"""
    import subprocess
    import re

    safe_name = escape_html(filename)

    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", doc_path],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout.strip()
        else:
            # fallback: 尝试直接读取
            with open(doc_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except Exception:
        return None

    parts = [_doc_html_header(safe_name)]

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            parts.append("<br>")
        elif len(line) < 60 and (line.endswith("：") or line.endswith(":") or
                                  "办法" in line or "规定" in line or "管理" in line):
            parts.append(f'<h3 class="doc-h3">{escape_html(line)}</h3>')
        else:
            parts.append(f'<p class="doc-p">{escape_html(line)}</p>')

    parts.append(_doc_html_footer())
    return "\n".join(parts)


def _excel_to_html(doc_path: str, filename: str) -> str:
    """将Excel文档转换为HTML预览表格"""
    fname_lower = filename.lower()

    safe_name = escape_html(filename)
    parts = [_doc_html_header(safe_name)]

    if fname_lower.endswith(".xls") and not fname_lower.endswith(".xlsx"):
        # 旧版.xls用xlrd
        try:
            import xlrd
            wb = xlrd.open_workbook(doc_path)
            for sheet_idx in range(wb.nsheets):
                sheet = wb.sheet_by_index(sheet_idx)
                parts.append(f'<h3 class="doc-h3">工作表: {escape_html(sheet.name)}</h3>')
                parts.append('<div class="table-wrapper"><table class="doc-table">')
                for row_idx in range(min(sheet.nrows, 500)):  # 最多500行
                    parts.append('<tr>')
                    for col_idx in range(sheet.ncols):
                        cell_value = str(sheet.cell_value(row_idx, col_idx))
                        parts.append(f'<td>{escape_html(cell_value)}</td>')
                    parts.append('</tr>')
                parts.append('</table></div>')
            parts.append(_doc_html_footer())
            return "\n".join(parts)
        except Exception:
            pass

    # .xlsx用openpyxl
    import openpyxl
    wb = openpyxl.load_workbook(doc_path, data_only=True)

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f'<h3 class="doc-h3">工作表: {escape_html(sheet_name)}</h3>')
        parts.append('<div class="table-wrapper"><table class="doc-table">')

        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            if row_count > 500:
                break
            parts.append('<tr>')
            for cell in row:
                cell_text = str(cell) if cell is not None else ""
                parts.append(f'<td>{escape_html(cell_text)}</td>')
            parts.append('</tr>')
            row_count += 1

        parts.append('</table></div>')

    wb.close()
    parts.append(_doc_html_footer())
    return "\n".join(parts)


def _doc_html_header(title: str) -> str:
    """生成文档HTML预览页面的头部"""
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei",
                         "Helvetica Neue", sans-serif;
            max-width: 960px; margin: 0 auto; padding: 24px 20px 60px;
            line-height: 1.8; color: #1a1a2e; background: #f8f9fb;
        }}
        .doc-header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: #fff; padding: 20px 28px; border-radius: 12px;
            margin-bottom: 28px; box-shadow: 0 4px 15px rgba(102,126,234,0.3);
        }}
        .doc-header h1 {{ font-size: 20px; font-weight: 600; }}
        .doc-header .doc-type {{ font-size: 12px; opacity: 0.85; margin-top: 4px; }}
        .back-link {{
            display: inline-block; margin-bottom: 20px; color: #667eea;
            text-decoration: none; font-size: 14px;
        }}
        .back-link:hover {{ text-decoration: underline; }}
        .doc-h2 {{
            font-size: 18px; font-weight: 600; color: #2d3436;
            margin: 24px 0 12px; padding-bottom: 8px;
            border-bottom: 2px solid #667eea;
        }}
        .doc-h3 {{
            font-size: 16px; font-weight: 600; color: #444;
            margin: 18px 0 8px;
        }}
        .doc-h4 {{
            font-size: 14px; font-weight: 600; color: #555;
            margin: 14px 0 6px;
        }}
        .doc-p {{
            margin: 6px 0; text-indent: 2em; font-size: 15px;
            color: #333; line-height: 1.9;
        }}
        .doc-table {{
            width: 100%; border-collapse: collapse; margin: 12px 0 20px;
            font-size: 13px; background: #fff; border-radius: 8px;
            overflow: hidden; box-shadow: 0 1px 4px rgba(0,0,0,0.06);
        }}
        .doc-table td {{
            border: 1px solid #e5e7eb; padding: 8px 12px;
            vertical-align: top; word-break: break-all;
        }}
        .doc-table tr:first-child td {{
            background: #f0f2ff; font-weight: 600; color: #444;
        }}
        .doc-table tr:hover td {{ background: #fafbff; }}
        .table-wrapper {{
            overflow-x: auto; -webkit-overflow-scrolling: touch;
        }}
        .doc-footer {{
            margin-top: 32px; padding-top: 16px;
            border-top: 1px solid #e5e7eb; text-align: center;
            font-size: 12px; color: #999;
        }}
    </style>
</head>
<body>
    <a class="back-link" href="/">← 返回对话</a>
    <div class="doc-header">
        <h1>📄 {title}</h1>
        <div class="doc-type">知识库文档预览</div>
    </div>
"""


def _doc_html_footer() -> str:
    """生成文档HTML预览页面的尾部"""
    return """
    <div class="doc-footer">
        <p>— 文档预览由智能问答系统提供 —</p>
    </div>
</body>
</html>"""
